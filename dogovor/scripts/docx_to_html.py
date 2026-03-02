#!/usr/bin/env python3
"""
Безопасный парсинг .docx: извлечение текста и структуры без потери форматирования.
Использует python-docx — не меняет исходный файл, только читает.

Использование:
  python docx_to_html.py путь/к/файлу.docx
  или
  python docx_to_html.py путь/к/файлу.docx --output вывод.html

Выводит в stdout или в файл HTML-подобную разметку параграфов и текста.
"""
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.text.paragraph import Paragraph
except ImportError:
    print("Установите python-docx: pip install python-docx", file=sys.stderr)
    sys.exit(1)


def _get_paragraph_html(p: "Paragraph") -> str:
    """Преобразует параграф в HTML-подобную строку с сохранением форматирования."""
    parts = []
    for run in p.runs:
        text = run.text
        if not text:
            continue
        if run.bold:
            text = f"<b>{text}</b>"
        if run.italic:
            text = f"<i>{text}</i>"
        if run.underline:
            text = f"<u>{text}</u>"
        parts.append(text)
    return "".join(parts)


def docx_to_html(docx_path: str) -> str:
    """
    Читает .docx и возвращает HTML-подобную строку.
    Только чтение — исходный файл не изменяется.
    """
    doc = Document(docx_path)
    out = []
    for p in doc.paragraphs:
        html = _get_paragraph_html(p)
        if html.strip():
            out.append(f"<p>{html}</p>")
        else:
            out.append("<p></p>")
    for table in doc.tables:
        out.append("<table>")
        for row in table.rows:
            out.append("<tr>")
            for cell in row.cells:
                out.append(f"<td>{cell.text}</td>")
            out.append("</tr>")
        out.append("</table>")
    return "\n".join(out)


def main():
    if len(sys.argv) < 2:
        print("Использование: python docx_to_html.py <file.docx> [--output file.html]", file=sys.stderr)
        sys.exit(1)
    path = Path(sys.argv[1])
    if not path.exists():
        print(f"Файл не найден: {path}", file=sys.stderr)
        sys.exit(1)
    html = docx_to_html(str(path))
    if "--output" in sys.argv:
        idx = sys.argv.index("--output")
        if idx + 1 < len(sys.argv):
            out_path = Path(sys.argv[idx + 1])
            out_path.write_text(html, encoding="utf-8")
            print(f"Сохранено в {out_path}", file=sys.stderr)
        else:
            print(html)
    else:
        print(html)


if __name__ == "__main__":
    main()
